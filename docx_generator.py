"""
DOCX Generator Module
Generates Word documents with preserved formatting from OCR data.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE


class DOCXGenerator:
    """Generates formatted Word documents from OCR data."""
    
    def __init__(self):
        self.default_font = 'Calibri'
        self.default_size = 11
    
    def create_document(self, formatted_data, output_path):
        """
        Create a Word document from formatted OCR data.
        
        Args:
            formatted_data: Dictionary containing OCR data with formatting
            output_path: Path to save the output document
        """
        doc = Document()
        
        # Set up document styles
        self._setup_document_styles(doc)
        
        # Group words by paragraph and line
        paragraphs_data = self._group_by_paragraphs(formatted_data)
        
        # Create paragraphs in document
        for par_data in paragraphs_data:
            self._add_paragraph_to_document(doc, par_data)
        
        # Save document
        doc.save(output_path)
    
    def _setup_document_styles(self, doc):
        """Set up default document styles."""
        style = doc.styles['Normal']
        font = style.font
        font.name = self.default_font
        font.size = Pt(self.default_size)
    
    def _group_by_paragraphs(self, formatted_data):
        """
        Group words by paragraph and line for document generation.
        
        Args:
            formatted_data: Dictionary containing OCR data with formatting
            
        Returns:
            list: List of paragraph data dictionaries
        """
        words = formatted_data.get('words', [])
        paragraphs_info = formatted_data.get('paragraphs', [])
        
        if not words:
            return []
        
        # Group words by paragraph number
        par_groups = {}
        
        for word in words:
            par_key = (word.get('block_num', 1), word.get('par_num', 1))
            
            if par_key not in par_groups:
                par_groups[par_key] = {
                    'words': [],
                    'lines': {},
                    'alignment': 'left'
                }
            
            par_groups[par_key]['words'].append(word)
            
            # Group by line within paragraph
            line_num = word.get('line_num', 1)
            if line_num not in par_groups[par_key]['lines']:
                par_groups[par_key]['lines'][line_num] = []
            
            par_groups[par_key]['lines'][line_num].append(word)
        
        # Add alignment info from paragraphs
        for i, par_key in enumerate(par_groups):
            if i < len(paragraphs_info):
                par_groups[par_key]['alignment'] = paragraphs_info[i].get('alignment', 'left')
        
        # Convert to list
        paragraphs_list = []
        for par_key in sorted(par_groups.keys()):
            par_data = par_groups[par_key]
            
            # Sort lines and words
            sorted_lines = []
            for line_num in sorted(par_data['lines'].keys()):
                line_words = par_data['lines'][line_num]
                # Sort words by x position (left to right)
                line_words.sort(key=lambda w: w['bbox']['x'])
                sorted_lines.append(line_words)
            
            paragraphs_list.append({
                'lines': sorted_lines,
                'alignment': par_data['alignment']
            })
        
        return paragraphs_list
    
    def _add_paragraph_to_document(self, doc, par_data):
        """
        Add a paragraph to the Word document.
        
        Args:
            doc: Document object
            par_data: Paragraph data dictionary
        """
        paragraph = doc.add_paragraph()
        
        # Set paragraph alignment
        alignment = par_data.get('alignment', 'left')
        if alignment == 'center':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment == 'right':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif alignment == 'justify':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Add words with formatting
        lines = par_data.get('lines', [])
        
        for i, line_words in enumerate(lines):
            # Add line break between lines (except first)
            if i > 0:
                paragraph.add_run().add_break()
            
            # Add words in this line
            for j, word in enumerate(line_words):
                run = paragraph.add_run(word['text'])
                
                # Apply formatting
                self._apply_formatting(run, word)
                
                # Add space after word (except last word)
                if j < len(line_words) - 1:
                    run = paragraph.add_run(" ")
    
    def _apply_formatting(self, run, word_data):
        """
        Apply formatting to a run.
        
        Args:
            run: Run object
            word_data: Dictionary containing word formatting info
        """
        # Set font properties
        font = run.font
        
        # Font name
        font.name = self.default_font
        
        # Font size
        font_size = word_data.get('font_size', self.default_size)
        font.size = Pt(font_size)
        
        # Bold
        if word_data.get('is_bold', False):
            font.bold = True
        
        # Italic
        if word_data.get('is_italic', False):
            font.italic = True
        
        # Text color (default black)
        text_color = word_data.get('text_color', 'black')
        if text_color == 'black':
            font.color.rgb = RGBColor(0, 0, 0)
        elif text_color == 'red':
            font.color.rgb = RGBColor(255, 0, 0)
        elif text_color == 'blue':
            font.color.rgb = RGBColor(0, 0, 255)
    
    def create_simple_document(self, text, output_path):
        """
        Create a simple document without formatting (fallback method).
        
        Args:
            text: Plain text content
            output_path: Path to save the output document
        """
        doc = Document()
        
        # Split text into paragraphs
        paragraphs = text.split('\n\n')
        
        for par_text in paragraphs:
            if par_text.strip():
                doc.add_paragraph(par_text.strip())
        
        doc.save(output_path)
